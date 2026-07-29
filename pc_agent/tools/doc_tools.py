from pathlib import Path
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

from pc_agent.config import DOCUMENTS_DIR, DEFAULT_WORKSPACE

def create_word_document(title: str, content_sections: list, output_filename: str = None) -> str:
    """
    Generate a formatted Microsoft Word (.docx) document.
    
    content_sections format:
    [
        {"heading": "Introduction", "body": "Text paragraph content here..."},
        {"heading": "Main Topics", "body": "Detailed notes..."}
    ]
    """
    if docx is None:
        return "Error: 'python-docx' library is not installed. Run 'pip install python-docx'."
        
    doc = docx.Document()
    
    # Title
    p_title = doc.add_heading(level=0)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("") # Spacing
    
    # Add Sections
    for sec in content_sections:
        heading = sec.get("heading", "")
        body = sec.get("body", "")
        
        if heading:
            h = doc.add_heading(heading, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x00, 0x55, 0x99)
            
        if body:
            p = doc.add_paragraph(body)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)
            
    if not output_filename:
        safe_title = "".join([c if c.isalnum() else "_" for c in title])[:30]
        output_filename = f"Assignment_{safe_title}.docx"
        
    save_path = DOCUMENTS_DIR / output_filename
    if not save_path.name.endswith(".docx"):
        save_path = save_path.with_suffix(".docx")
        
    doc.save(save_path)
    return str(save_path)
